from __future__ import annotations

from pathlib import Path

from ..errors import RAGValidationError
from ..schemas import LoadedDocument, TextBlock
from .canonical import serialize_canonical_blocks
from .normalizers import normalize_plain_text


class DocxFileLoader:
    loader_type = "docx"

    def __init__(self, *, loader_version: str) -> None:
        self.loader_version = loader_version

    def load(self, *, path: Path, source_name: str) -> LoadedDocument:
        try:
            import docx
        except ImportError as exc:
            raise RAGValidationError("docx parsing requires python-docx dependency") from exc

        document = docx.Document(str(path))
        blocks: list[TextBlock] = []
        for idx, paragraph in enumerate(document.paragraphs, start=1):
            text = normalize_plain_text(paragraph.text)
            if not text:
                continue
            blocks.append(
                TextBlock(
                    text=text,
                    metadata={"source": source_name, "section": f"paragraph-{idx}"},
                )
            )
        return LoadedDocument(
            loader_type=self.loader_type,
            loader_version=self.loader_version,
            extraction_method="docx_paragraphs",
            blocks=blocks,
            derived_text=serialize_canonical_blocks(blocks),
        )
